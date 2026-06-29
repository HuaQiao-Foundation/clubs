#!/usr/bin/env python3
"""Build RC Georgetown Project Summary Word forms — matched to the house style
of RCG_Project_Brief / Approval FINAL templates (Calibri, azure header block,
thin grid, azure section bars, tick-boxes).

Run with the shared monorepo venv (see docs/document-toolchain.md):
    /Users/randaleastman/dev/clubs/.venv/bin/python build-project-summary-forms.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- House palette (from FINAL templates) ----
AZURE    = "0067C8"
AZURE_DK = "005099"
GOLD     = "F7A81B"
LTBLUE   = "C8DCEE"
BODY     = "222222"
MUTED    = "777777"
GREY_BG  = "F5F5F5"
WHITE    = "FFFFFF"
BORDER   = "BFBFBF"   # thin grey grid (softer than pure black)
FONT     = "Calibri"

def _rgb(h): return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))

def el(tag, **attrs):
    e=OxmlElement(tag)
    for k,v in attrs.items(): e.set(qn(k.replace('_',':')), str(v))
    return e

def shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(el('w:shd', w_val='clear', w_color='auto', w_fill=fill))

def borders(cell, color=BORDER, sz=4):
    tcPr=cell._tc.get_or_add_tcPr(); b=el('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b.append(el(f'w:{edge}', w_val='single', w_sz=sz, w_space='0', w_color=color))
    tcPr.append(b)

def margins(cell, t=70, b=70, l=110, r=110):
    tcPr=cell._tc.get_or_add_tcPr(); m=el('w:tcMar')
    for edge,val in (('top',t),('bottom',b),('left',l),('right',r)):
        m.append(el(f'w:{edge}', w_w=val, w_type='dxa'))
    tcPr.append(m)

def row_height(row, tw):
    row._tr.get_or_add_trPr().append(el('w:trHeight', w_val=tw, w_hRule='atLeast'))

def fixed_layout(table, col_inches):
    """Force Word to honor explicit column widths: fixed layout + total width + a tblGrid.
    Without this, Word auto-fits columns and ignores per-cell .width (the cause of the
    over-wide column 1 in the SDG table)."""
    tw=sum(col_inches); twips=[int(round(w*1440)) for w in col_inches]
    tblPr=table._tbl.tblPr
    # fixed layout
    lay=tblPr.find(qn('w:tblLayout'))
    if lay is None: lay=el('w:tblLayout'); tblPr.append(lay)
    lay.set(qn('w:type'),'fixed')
    # total table width
    w=tblPr.find(qn('w:tblW'))
    if w is None: w=el('w:tblW'); tblPr.append(w)
    w.set(qn('w:w'), str(int(round(tw*1440)))); w.set(qn('w:type'),'dxa')
    # explicit grid
    old=table._tbl.find(qn('w:tblGrid'))
    if old is not None: table._tbl.remove(old)
    grid=el('w:tblGrid')
    for tp in twips: grid.append(el('w:gridCol', w_w=tp))
    table._tbl.insert(list(table._tbl).index(tblPr)+1, grid)
    # and set every cell width to match (belt and suspenders)
    for r in table.rows:
        for i,c in enumerate(r.cells):
            c.width=Inches(col_inches[i])

def run(p, text, *, size=9.5, bold=False, italic=False, color=BODY, font=FONT):
    r=p.add_run(text); r.font.name=font; r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.color.rgb=_rgb(color)
    rpr=r._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=el('w:rFonts'); rpr.insert(0,rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),font)
    return r

def para(container, *, before=0, after=0, line=1.12, align=None):
    p=container.add_paragraph(); pf=p.paragraph_format
    pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line
    if align: p.alignment=align
    return p

def cell_para(cell):
    p=cell.paragraphs[0]; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    return p

# ---------- header block (azure bar) ----------
def header_block(doc, subtitle):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.allow_autofit=False
    c=t.rows[0].cells[0]; c.width=Inches(6.7)
    shade(c, AZURE); borders(c, AZURE, 4); margins(c, t=90,b=90,l=160,r=160); row_height(t.rows[0], 900)
    p=cell_para(c)
    run(p,"ROTARY CLUB OF GEORGETOWN", size=11, bold=True, color=GOLD)
    p2=cell.add_paragraph() if False else c.add_paragraph(); p2.paragraph_format.space_before=Pt(2); p2.paragraph_format.space_after=Pt(0)
    run(p2,"Project Summary Form", size=18, bold=True, color=WHITE)
    p3=c.add_paragraph(); p3.paragraph_format.space_before=Pt(1); p3.paragraph_format.space_after=Pt(0)
    run(p3, subtitle, size=9, color=LTBLUE)
    para(doc, after=4)  # gap after header

# ---------- section bar (canonical: darker azure, uppercase) ----------
def section(doc, title):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.allow_autofit=False
    c=t.rows[0].cells[0]; c.width=Inches(6.7)
    shade(c, AZURE_DK); borders(c, AZURE_DK, 4); margins(c, t=40,b=40,l=110,r=110); row_height(t.rows[0], 300)
    p=cell_para(c); run(p, title.upper(), size=10, bold=True, color=WHITE)
    para(doc, after=0, before=0)  # tight gap; table spacing handles it

# ---------- reference appendix table (azure-dk header row + grid body) ----------
def appendix_table(doc, headers, rows, widths):
    t=doc.add_table(rows=0, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.allow_autofit=False
    hr=t.add_row(); row_height(hr, 300)
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.width=Inches(widths[i]); shade(c, AZURE_DK); borders(c); margins(c)
        run(cell_para(c), h, size=10, bold=True, color=WHITE)
    for row in rows:
        rr=t.add_row(); row_height(rr, 280)
        for i,val in enumerate(row):
            c=rr.cells[i]; c.width=Inches(widths[i]); borders(c); margins(c)
            run(cell_para(c), val, size=10, color=BODY)
    fixed_layout(t, widths)   # make Word honor the column ratio
    para(doc, after=6)

# ---------- field rows (label | answer) ----------
def fields(doc, rows, label_in=2.5, ans_in=4.2):
    t=doc.add_table(rows=0, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.allow_autofit=False
    for label, hint, answer, tall in rows:
        r=t.add_row(); row_height(r, 620 if tall else 380)
        lc, ac = r.cells
        lc.width=Inches(label_in); ac.width=Inches(ans_in)
        shade(lc, GREY_BG)
        for c in (lc,ac): borders(c); margins(c)
        lp=cell_para(lc); run(lp, label, size=10, bold=True, color=BODY)
        if hint: run(lp, "  "+hint, size=8, italic=True, color=MUTED)
        ap=cell_para(ac)
        if answer:
            lines=answer.split("\n")
            r=run(ap, lines[0], size=9.5, color=BODY)
            for extra in lines[1:]:
                r.add_break()                       # soft line break within the cell
                run(ap, extra, size=9.5, color=BODY)
    fixed_layout(t, [label_in, ans_in])             # honor label/answer column widths
    para(doc, after=6)  # space between sections

def F(label, answer="", hint=None, tall=False): return (label, hint, answer, tall)

def build(path, *, draft_for=None, prefill=None, extras_hint=None):
    pf=prefill or {}
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9)
    st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(9.5); st.font.color.rgb=_rgb(BODY)

    subtitle = (f"{pf.get('Project name','')} — draft for {draft_for}" if draft_for
                else "Record of a club service project")
    header_block(doc, subtitle)

    first_name = draft_for.split()[0] if draft_for else None
    # (No greeting/intro paragraph in the form body — per CEO, removed 2026-06-29.)

    g=lambda k: pf.get(k,"")

    section(doc,"Project Identity")
    fields(doc, [
        F("Project name", g("Project name")),
        F("Rotary year", g("Rotary year"), "e.g. 2024-25"),
        F("Project champion", g("Champion"), "member who led it"),
    ])
    section(doc,"The Story")
    fields(doc, [
        F("What was it?", g("What was it?"), "one or two sentences", tall=True),
        F("Why did the club do it?", g("Why"), "the need or opportunity it addressed", tall=True),
        F("What was the impact?", g("Impact"), "what changed — in human terms", tall=True),
    ])
    section(doc,"Classification")
    fields(doc, [
        F("Area of Focus", g("Area of Focus"), "see appendix — pick one of the seven"),
        F("UN SDG", g("UN SDG"), "see appendix — e.g. SDG 6 Clean Water & Sanitation"),
        F("Funding source", g("Funding"), "Club · District Grant · Global Grant · Joint"),
        F("Project status", g("Status"), "Planning · Execution · Completed · Dropped"),
    ])
    section(doc,"When & Where")
    fields(doc, [
        F("Start date", g("Start date"), "e.g. 2024-09"),
        F("Completion date", g("Completion date"), "blank if ongoing"),
        F("Location", g("Location")),
    ])
    section(doc,"By the Numbers")
    fields(doc, [
        F("Beneficiaries", g("Beneficiaries"), "number of people helped"),
        F("Project value (RM)", g("Value"), "total cost or value"),
        F("Volunteer hours", g("Volunteer hours"), "optional — estimate is fine"),
    ])
    section(doc,"Partners")
    fields(doc, [
        F("Partner organisations", g("Partners"), "sponsors, co-clubs, NGOs, or “none”", tall=True),
    ])
    section(doc,"Sustainability & Community Participation")
    fields(doc, [
        F("Sustainability measures", g("Sustainability"), "how will benefits continue after Rotary's involvement ends?", tall=True),
        F("Community participation", g("Community participation"), "how did the target community contribute to or have a stake in this?", tall=True),
    ])
    section(doc,"Extras  (all optional)")
    fields(doc, [
        F("Photos available?", g("Photos"), "yes / no — and where to find them"),
        F("Publicity", g("Publicity"), "links to articles, social media, write-ups"),
        F("Lessons learned", g("Lessons"), "would you do it again? what would you change?", tall=True),
        F("Anything else?", g("Anything else"), extras_hint or "anything the club should record", tall=True),
    ])
    fp=para(doc, before=10)
    run(fp, "Thank you. This helps us build our Service Projects archive and keeps all members — "
            "new and long-standing — familiar with our club's accomplishments.",
        size=8.5, italic=True, color=MUTED)

    # ---- Reference appendix (Areas of Focus + UN SDGs) ----
    para(doc, before=8)
    rp=para(doc, after=2); run(rp, "REFERENCE APPENDIX", size=11, bold=True, color=AZURE_DK)
    rp2=para(doc, after=6); run(rp2, "Use these tables when filling in the Classification section above.", size=9, italic=True, color=MUTED)
    sp=para(doc, after=2); run(sp, "Rotary Areas of Focus", size=10, bold=True, color=BODY)
    appendix_table(doc,
        ["Area of Focus","Typical activities"],
        [["Peacebuilding & Conflict Prevention","Training, mediation, refugee support"],
         ["Disease Prevention & Treatment","Health infrastructure, immunisation, clean clinics"],
         ["Water, Sanitation & Hygiene","Wells, latrines, water filters, WASH education"],
         ["Maternal & Child Health","Pre-natal care, nutrition, infant survival"],
         ["Basic Education & Literacy","Schools, teacher training, reading programs"],
         ["Community Economic Development","Microfinance, vocational training, livelihoods"],
         ["Supporting the Environment","Reforestation, waste management, sustainability"]],
        [2.7,4.0])
    sp2=para(doc, after=2); run(sp2, "United Nations Sustainable Development Goals (SDGs)", size=10, bold=True, color=BODY)
    appendix_table(doc,
        ["#","Goal","Description"],
        [["SDG 1","No Poverty","End poverty in all its forms"],
         ["SDG 2","Zero Hunger","Food security and better nutrition"],
         ["SDG 3","Good Health","Healthy lives and well-being for all ages"],
         ["SDG 4","Quality Education","Access to good education and lifelong learning"],
         ["SDG 5","Gender Equality","Empower all women and girls"],
         ["SDG 6","Clean Water & Sanitation","Safe water and sanitation for all"],
         ["SDG 7","Clean Energy","Affordable, reliable, modern energy"],
         ["SDG 8","Decent Work","Economic growth and fair employment"],
         ["SDG 9","Innovation & Infrastructure","Resilient infrastructure and industry"],
         ["SDG 10","Reduced Inequalities","Less inequality within and between countries"],
         ["SDG 11","Sustainable Cities","Safe, resilient, well-planned communities"],
         ["SDG 12","Responsible Consumption","Produce and consume sustainably"],
         ["SDG 13","Climate Action","Act on climate change and its impacts"],
         ["SDG 14","Life Below Water","Protect oceans and marine life"],
         ["SDG 15","Life on Land","Protect forests, land, and biodiversity"],
         ["SDG 16","Peace & Justice","Peaceful societies and strong institutions"],
         ["SDG 17","Partnerships","Work together to achieve the goals"]],
        [1.12, 2.23, 3.35])   # narrow # : medium Goal : wide Description = 1 : 2 : 3
    doc.save(path); print("wrote", path)

# All generated .docx land in the gitignored forms/ folder (artifacts, not source).
OUT="/Users/randaleastman/dev/clubs/apps/georgetown/forms"

# NOTE: this script does NOT generate the blank template
# `RC-Georgetown-Project-Summary-Form.docx` — that file is the CEO's canonical,
# hand-made (Claude Console) template and must never be overwritten by this script.
# Likewise `Project-Summary-Aquaponics-Workshop-Mike.docx` (no "for") is CEO-authored.
# This script only emits the PRE-FILLED project summaries below.

build(f"{OUT}/Project-Summary-Aquaponics-Workshop-for-Mike.docx", draft_for="Mike Jackman",
      extras_hint="the club ran an earlier aquaponics installation at St Nicholas Home for the Blind in 2021-22 — was this related?",
      prefill={"Project name":"Aquaponics Workshop","Rotary year":"2024-25","Champion":"Mike Jackman",
               "Area of Focus":"Supporting the Environment","UN SDG":"SDG 12 — Responsible Consumption",
               "Funding":"Club","Status":"Completed","Location":"Penang"})

# Shrijan Maram Pitchmasters Scholarship — completed record (DB project 96ceba73-…)
build(f"{OUT}/Project-Summary-Shrijan-Maram-Scholarship.docx",
      prefill={
        "Project name":"Shrijan Maram Pitchmasters Scholarship",
        "Rotary year":"2025-26",
        "Champion":"Randal Eric Eastman",
        "What was it?":("A project to sponsor the Pitchmasters charter dues (RM 388) for an 18-year-old "
                        "student, Shrijan Maram — giving a proven young local the communication training "
                        "and platform to grow his impact."),
        "Why":("Shrijan is a capable, driven student who already has the skills and track record but "
               "lacked the communication infrastructure to take his work further. For a small sum, the "
               "club could remove the cost barrier to that training."),
        "Impact":("“Shrijan Maram is not potential — he is proven. He has the skills. He has the drive. "
                  "He has the track record. He needs the communication infrastructure to take his impact "
                  "global. For USD 88, Rotary Club of Georgetown can help a local prodigy develop the voice "
                  "to match his vision. This isn’t sponsorship — it’s partnership.”"),
        "Area of Focus":"Basic Education & Literacy",
        "UN SDG":"SDG 4 — Quality Education",
        "Funding":"Club",
        "Status":"Completed",
        "Start date":"2025-10-01",
        "Completion date":"2025-10-06",
        "Location":"Penang",
        "Beneficiaries":"1 (Shrijan Maram)",
        "Value":"RM 388.00 (~USD 88)",
        "Volunteer hours":"—",
        "Partners":("PG GAMA Supermarket & Department Store — funding partner; corporate sponsor who "
                    "underwrote the RM 388 charter dues (contact: Andrew TK Lim).\n"
                    "Pitchmasters — program partner; the public-speaking club Shrijan joins via the "
                    "sponsored charter dues."),
        "Photos":"Yes — project image on file (Supabase project-images/96ceba73-…jpg).",
        "Publicity":"— (none recorded)",
        "Lessons":"Would repeat: Yes. No lessons-learned note recorded.",
        "Anything else":("Possible follow-ups: confirm whether this is a one-off or the start of a recurring "
                         "Pitchmasters scholarship; track Shrijan’s progress in Pitchmasters as an outcome "
                         "data point."),
      })

# Christmas Orphan Care Project 2024 — completed record (DB project 463bbd9f-…)
build(f"{OUT}/Project-Summary-Christmas-Orphan-Care-2024.docx",
      prefill={
        "Project name":"Christmas Orphan Care Project",
        "Rotary year":"2024-25",
        "Champion":"Yew-Aun Soh",
        "What was it?":"The club's 3rd Annual Christmas celebration and gift distribution for a local orphanage.",
        "Why":("A now-annual commitment: bringing a Christmas celebration and gifts to children in a local "
               "orphanage who would otherwise go without."),
        "Impact":"3 community organizations helping 300 families (record figure); 20 children directly celebrated and received gifts.",
        "Area of Focus":"Maternal & Child Health",
        "UN SDG":"SDG 3 — Good Health & Well-being",
        "Funding":"Club",
        "Status":"Completed",
        "Start date":"2024-12-14",
        "Completion date":"2024-12-14 (single-day event)",
        "Location":"Georgetown, Penang",
        "Beneficiaries":"20 (children at the orphanage; wider note: 300 families via 3 community organizations)",
        "Value":"RM 4,999.00",
        "Volunteer hours":"—",
        "Partners":"None recorded. The impact note references “3 community organizations” — if these were formal partners, name them and we can add them to the record.",
        "Photos":"Yes — project image on file (Supabase project-images/463bbd9f-…jpg).",
        "Publicity":"— (none recorded)",
        "Lessons":"Would repeat: Yes. No lessons-learned note recorded.",
        "Anything else":("This is the 3rd Annual edition — part of a recurring Christmas series. "
                         "Worth naming the orphanage and the 3 community organizations for the record."),
      })

# Christmas Orphan Care Project 2025 — "4th Annual" DRAFT skeleton for Yew-Aun Soh to complete
# (no 2025 DB record yet; recurring fields carried forward, year-specific fields left blank)
build(f"{OUT}/Project-Summary-Christmas-Orphan-Care-2025-DRAFT.docx", draft_for="Yew-Aun Soh",
      extras_hint=("this is the recurring annual Christmas series — the 2021, 2022, and 2024 editions are known; "
                   "was there a 2023 edition? Naming the orphanage would help the record."),
      prefill={
        "Project name":"Christmas Orphan Care Project  (4th Annual — adjust if changed)",
        "Rotary year":"2025-26",
        "Champion":"Yew-Aun Soh",
        "Area of Focus":"Maternal & Child Health  (carried — change if you disagree)",
        "UN SDG":"SDG 3 — Good Health  (carried — confirm)",
        "Funding":"Club  (carried — confirm)",
        "Location":"Georgetown, Penang  (carried — confirm venue / orphanage)",
        # Year-specific fields intentionally blank: What was it?, Why, Impact, Status,
        # Start/Completion date, Beneficiaries, Value, Volunteer hours, Partners, Photos, etc.
      })
