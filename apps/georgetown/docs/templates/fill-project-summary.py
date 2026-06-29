#!/usr/bin/env python3
"""Fill the CEO's canonical RC Georgetown Project Summary template with real data.

This does NOT build a form from scratch — it OPENS the canonical template
(docs/templates/RC-Georgetown-Project-Summary-Form.docx),
finds each field by its label, and replaces the placeholder hint in the
answer cell with the project's value. All of the CEO's styling, section
bars, layout, and the reference appendix are preserved exactly.

Run with the shared venv:
    /Users/randaleastman/dev/clubs/.venv/bin/python fill-project-summary.py
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

TEMPLATE = "/Users/randaleastman/dev/clubs/apps/georgetown/docs/templates/RC-Georgetown-Project-Summary-Form.docx"
OUT      = "/Users/randaleastman/dev/clubs/apps/georgetown/forms"
BODY     = RGBColor(0x22,0x22,0x22)

# Field tables in the template are 1x2: [label | answer-placeholder].
# We match on the label (left cell) and rewrite the right cell.

def set_answer(cell, text, *, draft=False):
    """Replace a field cell's content with `text`, matching the template's
    answer-cell font (Calibri ~11pt). Empty text = leave the placeholder."""
    if text is None or text == "":
        return
    # wipe existing paragraphs/runs but keep paragraph formatting of the first
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # keep rows snug: clamp any large trailing space (the template's 30pt story
    # fields) down to 3pt so filled cells don't carry a tall empty gap below the text
    sa = p.paragraph_format.space_after
    if sa is None or sa.pt > 3:
        p.paragraph_format.space_after = Pt(3)
    # support multi-line via \n (soft breaks)
    lines = str(text).split("\n")
    r = p.add_run(lines[0])
    for ln in lines[1:]:
        r.add_break(); r = p.add_run(ln)
    for run in p.runs:
        run.font.name = "Calibri"; run.font.size = Pt(11)
        run.font.color.rgb = BODY
        run.italic = draft   # draft placeholders rendered italic to read as "to fill"
        rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = rpr.makeelement(qn('w:rFonts'), {}); rpr.insert(0, rf)
        for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), "Calibri")

def fill(template, out, answers, *, draft_fields=()):
    """answers: {label: value}. draft_fields: labels whose value is a
    'please fill' placeholder (rendered italic)."""
    doc = Document(template)
    matched = set()
    for t in doc.tables:
        if len(t.columns) != 2 or len(t.rows) != 1:
            continue
        label = t.rows[0].cells[0].text.strip()
        # tolerate trailing punctuation / case
        key = next((k for k in answers if k.lower() == label.lower()), None)
        if key is not None:
            set_answer(t.rows[0].cells[1], answers[key], draft=(key in draft_fields))
            matched.add(key)
    missing = set(answers) - matched
    if missing:
        print(f"  WARNING — labels not found in template: {sorted(missing)}")
    doc.save(out)
    print("wrote", out)

# ---------------------------------------------------------------------------
# Shrijan Maram Pitchmasters Scholarship — completed record (DB 96ceba73-…)
fill(TEMPLATE, f"{OUT}/Project-Summary-Shrijan-Maram-Scholarship.docx", {
    "Project name":"Shrijan Maram Pitchmasters Scholarship",
    "Rotary year":"2025-26",
    "Project champion":"Randal Eric Eastman",
    "What was it?":("A project to sponsor the Pitchmasters charter dues (RM 388) for an 18-year-old "
                    "student, Shrijan Maram — giving a proven young local the communication training and "
                    "platform to grow his impact."),
    "Why did the club do it?":("Shrijan is a capable, driven student who already has the skills and track "
                               "record but lacked the communication infrastructure to take his work further. "
                               "For a small sum, the club could remove the cost barrier to that training."),
    "What was the impact?":("“Shrijan Maram is not potential — he is proven. He has the skills, the drive, the "
                            "track record. He needs the communication infrastructure to take his impact global. "
                            "For USD 88, Rotary Club of Georgetown can help a local prodigy develop the voice to "
                            "match his vision. This isn’t sponsorship — it’s partnership.”"),
    "Area of Focus":"Basic Education & Literacy",
    "UN SDG":"SDG 4 — Quality Education",
    "Funding source":"Club",
    "Project status":"Completed",
    "Start date":"2025-10-01",
    "Completion date":"2025-10-06",
    "Location":"Penang",
    "Beneficiaries":"1 (Shrijan Maram)",
    "Project value (RM)":"RM 388.00 (~USD 88)",
    "Volunteer hours":"—",
    "Partner organisations":("PG GAMA Supermarket & Department Store — funding partner; corporate sponsor "
                             "who underwrote the RM 388 dues (contact: Andrew TK Lim).\n"
                             "Pitchmasters — program partner; the public-speaking club Shrijan joins via the "
                             "sponsored charter dues."),
    "Sustainability measures":("The communication training Shrijan gains is a durable personal skill he carries "
                               "forward independently — the benefit continues with no further club input."),
    "Community participation":"—",
    "Photos available?":"Yes — project image on file (Supabase project-images/96ceba73-…jpg).",
    "Publicity":"—",
    "Lessons learned":"Would repeat: Yes.",
    "Anything else?":("Possible follow-up: confirm whether this is a one-off or the start of a recurring "
                      "Pitchmasters scholarship; track Shrijan’s progress as an outcome data point."),
})

# ---------------------------------------------------------------------------
# Christmas Orphan Care Project 2024 — completed record (DB 463bbd9f-…)
fill(TEMPLATE, f"{OUT}/Project-Summary-Christmas-Orphan-Care-2024.docx", {
    "Project name":"Christmas Orphan Care Project",
    "Rotary year":"2024-25",
    "Project champion":"Yew-Aun Soh",
    "What was it?":"The club's 3rd Annual Christmas celebration and gift distribution for a local orphanage.",
    "Why did the club do it?":("A now-annual commitment: bringing a Christmas celebration and gifts to children "
                               "in a local orphanage who would otherwise go without."),
    "What was the impact?":("3 community organizations helping 300 families (record figure); 20 children directly "
                            "celebrated and received gifts."),
    "Area of Focus":"Maternal & Child Health",
    "UN SDG":"SDG 3 — Good Health & Well-being",
    "Funding source":"Club",
    "Project status":"Completed",
    "Start date":"2024-12-14",
    "Completion date":"2024-12-14 (single-day event)",
    "Location":"Georgetown, Penang",
    "Beneficiaries":"20 (children at the orphanage; wider note: 300 families via 3 community organizations)",
    "Project value (RM)":"RM 4,999.00",
    "Volunteer hours":"—",
    "Partner organisations":("None recorded. The impact note references “3 community organizations” — if these "
                             "were formal partners, name them and we can add them to the record."),
    "Sustainability measures":("Recurring annual commitment — the project's continuation year on year is itself a "
                               "form of sustained engagement with the orphanage."),
    "Community participation":"—",
    "Photos available?":"Yes — project image on file (Supabase project-images/463bbd9f-…jpg).",
    "Publicity":"—",
    "Lessons learned":"Would repeat: Yes.",
    "Anything else?":("This is the 3rd Annual edition of a recurring Christmas series. Worth naming the orphanage "
                      "and the 3 community organizations for the record."),
})

# ---------------------------------------------------------------------------
# Christmas Orphan Care Project 2025 — "4th Annual", for Yew-Aun Soh to complete.
# No 2025 DB record yet: recurring fields carried; year-specific fields are italic
# "to fill" placeholders.
fill(TEMPLATE, f"{OUT}/Project-Summary-Christmas-Orphan-Care-2025.docx", {
    "Project name":"Christmas Orphan Care Project  (4th Annual — adjust if changed)",
    "Rotary year":"2025-26",
    "Project champion":"Yew-Aun Soh",
    "What was it?":"⟶ confirm or describe this year's edition (e.g. the 4th Annual Christmas celebration & gift distribution)",
    "Why did the club do it?":"⟶ carried: an annual commitment bringing Christmas and gifts to orphanage children — update if it shifted",
    "What was the impact?":"⟶ this year's numbers: children celebrated, gifts given, families/organizations reached",
    "Area of Focus":"Maternal & Child Health  (carried — change if you disagree)",
    "UN SDG":"SDG 3 — Good Health  (carried — confirm)",
    "Funding source":"Club  (carried — confirm)",
    "Project status":"⟶ Planning / Execution / Completed",
    "Start date":"⟶ e.g. 2025-12",
    "Completion date":"⟶",
    "Location":"Georgetown, Penang  (carried — confirm venue / orphanage)",
    "Beneficiaries":"⟶ 2024 was 20 children — this year's count",
    "Project value (RM)":"⟶ 2024 was RM 4,999 — this year's spend",
    "Volunteer hours":"⟶",
    "Partner organisations":"⟶ orphanage, community organizations, or sponsors this year — “none” is fine",
    "Sustainability measures":"⟶",
    "Community participation":"⟶",
    "Photos available?":"⟶",
    "Publicity":"⟶",
    "Lessons learned":"⟶",
    "Anything else?":"⟶ this is the recurring annual series (2021, 2022, 2024 known — was there a 2023?); naming the orphanage would help",
}, draft_fields={
    "What was it?","Why did the club do it?","What was the impact?","Project status","Start date",
    "Completion date","Beneficiaries","Project value (RM)","Volunteer hours","Partner organisations",
    "Sustainability measures","Community participation","Photos available?","Publicity","Lessons learned",
    "Anything else?",
})
